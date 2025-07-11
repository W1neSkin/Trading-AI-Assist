from fastapi import FastAPI, HTTPException, UploadFile, File, Depends, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any, Union
from datetime import datetime
from enum import Enum
import uuid
import asyncio
import logging
import io
import json
import hashlib
from contextlib import asynccontextmanager
import aiofiles
import tempfile
import os

# Document processing imports
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import pytesseract
from PIL import Image
import fitz  # PyMuPDF
import mammoth

# Local imports
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))

from shared.config import settings
from shared.database import postgresql_manager, mongodb_manager
from shared.messaging import hybrid_messaging_manager

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Enums
class DocumentType(str, Enum):
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    TXT = "txt"
    HTML = "html"
    RTF = "rtf"
    IMAGE = "image"

class ProcessingStatus(str, Enum):
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"

class ContentType(str, Enum):
    TEXT = "text"
    TABLE = "table"
    IMAGE = "image"
    METADATA = "metadata"

# Pydantic models
class DocumentInfo(BaseModel):
    id: str
    user_id: str
    filename: str
    original_filename: str
    file_size: int
    document_type: DocumentType
    content_hash: str
    processing_status: ProcessingStatus
    extracted_text: Optional[str]
    metadata: Dict[str, Any]
    created_at: datetime
    updated_at: datetime

class ProcessingResult(BaseModel):
    document_id: str
    status: ProcessingStatus
    extracted_content: Dict[str, Any]
    error_message: Optional[str]
    processing_time: float

class DocumentSearch(BaseModel):
    query: str
    document_types: Optional[List[DocumentType]] = None
    user_id: Optional[str] = None
    limit: int = 50

class ExtractedContent(BaseModel):
    text: str
    tables: List[Dict[str, Any]] = []
    images: List[Dict[str, Any]] = []
    metadata: Dict[str, Any] = {}

# Document Processor Classes
class PDFProcessor:
    async def extract_content(self, file_path: str) -> ExtractedContent:
        """Extract content from PDF"""
        try:
            extracted_text = ""
            tables = []
            images = []
            metadata = {}
            
            # Use pdfplumber for text and table extraction
            with pdfplumber.open(file_path) as pdf:
                metadata = {
                    "pages": len(pdf.pages),
                    "title": pdf.metadata.get('Title', ''),
                    "author": pdf.metadata.get('Author', ''),
                    "creator": pdf.metadata.get('Creator', ''),
                    "creation_date": str(pdf.metadata.get('CreationDate', '')),
                }
                
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    for table_num, table in enumerate(page_tables):
                        if table:
                            tables.append({
                                "page": page_num + 1,
                                "table_id": f"table_{page_num + 1}_{table_num + 1}",
                                "data": table,
                                "rows": len(table),
                                "columns": len(table[0]) if table else 0
                            })
            
            # Use PyMuPDF for image extraction
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc[page_num]
                image_list = page.get_images()
                
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    
                    if pix.n - pix.alpha < 4:  # GRAY or RGB
                        images.append({
                            "page": page_num + 1,
                            "image_id": f"img_{page_num + 1}_{img_index + 1}",
                            "width": pix.width,
                            "height": pix.height,
                            "colorspace": pix.colorspace.name if pix.colorspace else "unknown",
                            "size": len(pix.tobytes())
                        })
                    
                    pix = None
            
            doc.close()
            
            return ExtractedContent(
                text=extracted_text.strip(),
                tables=tables,
                images=images,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            raise

class DOCXProcessor:
    async def extract_content(self, file_path: str) -> ExtractedContent:
        """Extract content from DOCX"""
        try:
            # Extract text using python-docx
            doc = DocxDocument(file_path)
            
            extracted_text = ""
            tables = []
            metadata = {}
            
            # Extract document properties
            props = doc.core_properties
            metadata = {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables)
            }
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    extracted_text += para.text + "\n"
            
            # Extract tables
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                tables.append({
                    "table_id": f"table_{table_num + 1}",
                    "data": table_data,
                    "rows": len(table_data),
                    "columns": len(table_data[0]) if table_data else 0
                })
            
            # Alternative: Use mammoth for better HTML conversion
            try:
                with open(file_path, "rb") as docx_file:
                    result = mammoth.extract_raw_text(docx_file)
                    if result.value and len(result.value) > len(extracted_text):
                        extracted_text = result.value
            except Exception as e:
                logger.warning(f"Mammoth extraction failed: {e}")
            
            return ExtractedContent(
                text=extracted_text.strip(),
                tables=tables,
                images=[],  # Image extraction from DOCX requires additional processing
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            raise

class HTMLProcessor:
    async def extract_content(self, file_path: str) -> ExtractedContent:
        """Extract content from HTML"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract text
            extracted_text = soup.get_text(separator='\n', strip=True)
            
            # Extract tables
            tables = []
            for table_num, table in enumerate(soup.find_all('table')):
                table_data = []
                for row in table.find_all('tr'):
                    row_data = []
                    for cell in row.find_all(['td', 'th']):
                        row_data.append(cell.get_text(strip=True))
                    if row_data:
                        table_data.append(row_data)
                
                if table_data:
                    tables.append({
                        "table_id": f"table_{table_num + 1}",
                        "data": table_data,
                        "rows": len(table_data),
                        "columns": len(table_data[0]) if table_data else 0
                    })
            
            # Extract metadata
            metadata = {
                "title": soup.title.string if soup.title else "",
                "meta_description": "",
                "meta_keywords": "",
                "links": len(soup.find_all('a')),
                "images": len(soup.find_all('img')),
                "tables": len(tables)
            }
            
            # Extract meta tags
            for meta in soup.find_all('meta'):
                name = meta.get('name', '').lower()
                content = meta.get('content', '')
                if name == 'description':
                    metadata['meta_description'] = content
                elif name == 'keywords':
                    metadata['meta_keywords'] = content
            
            return ExtractedContent(
                text=extracted_text,
                tables=tables,
                images=[],
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error processing HTML: {e}")
            raise

class TXTProcessor:
    async def extract_content(self, file_path: str) -> ExtractedContent:
        """Extract content from TXT"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text_content = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("Could not decode text file with any standard encoding")
            
            # Basic text analysis
            lines = text_content.split('\n')
            words = text_content.split()
            
            metadata = {
                "lines": len(lines),
                "words": len(words),
                "characters": len(text_content),
                "encoding": encoding
            }
            
            return ExtractedContent(
                text=text_content,
                tables=[],
                images=[],
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error processing TXT: {e}")
            raise

class ImageProcessor:
    async def extract_content(self, file_path: str) -> ExtractedContent:
        """Extract content from images using OCR"""
        try:
            # Open image
            image = Image.open(file_path)
            
            # Extract text using OCR
            extracted_text = pytesseract.image_to_string(image)
            
            # Get image metadata
            metadata = {
                "width": image.width,
                "height": image.height,
                "format": image.format,
                "mode": image.mode,
                "size_bytes": os.path.getsize(file_path)
            }
            
            # Try to get EXIF data
            try:
                exif_data = image._getexif()
                if exif_data:
                    metadata["exif"] = {str(k): str(v) for k, v in exif_data.items()}
            except Exception:
                pass
            
            return ExtractedContent(
                text=extracted_text.strip(),
                tables=[],
                images=[{
                    "image_id": "main_image",
                    "width": image.width,
                    "height": image.height,
                    "format": image.format
                }],
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error processing image: {e}")
            raise

# Document Manager
class DocumentManager:
    def __init__(self):
        self.processors = {
            DocumentType.PDF: PDFProcessor(),
            DocumentType.DOCX: DOCXProcessor(),
            DocumentType.DOC: DOCXProcessor(),  # Use same processor
            DocumentType.HTML: HTMLProcessor(),
            DocumentType.TXT: TXTProcessor(),
            DocumentType.IMAGE: ImageProcessor()
        }
        self.upload_dir = "/tmp/document_uploads"
        os.makedirs(self.upload_dir, exist_ok=True)

    async def upload_document(self, user_id: str, file: UploadFile) -> DocumentInfo:
        """Upload and process document"""
        try:
            # Generate document ID
            document_id = str(uuid.uuid4())
            
            # Determine document type
            file_extension = file.filename.split('.')[-1].lower()
            document_type = self._get_document_type(file_extension)
            
            # Calculate file hash
            file_content = await file.read()
            content_hash = hashlib.sha256(file_content).hexdigest()
            
            # Save file temporarily
            temp_file_path = os.path.join(self.upload_dir, f"{document_id}.{file_extension}")
            
            async with aiofiles.open(temp_file_path, 'wb') as f:
                await f.write(file_content)
            
            # Create document record
            document_info = DocumentInfo(
                id=document_id,
                user_id=user_id,
                filename=f"{document_id}.{file_extension}",
                original_filename=file.filename,
                file_size=len(file_content),
                document_type=document_type,
                content_hash=content_hash,
                processing_status=ProcessingStatus.PENDING,
                extracted_text=None,
                metadata={},
                created_at=datetime.utcnow(),
                updated_at=datetime.utcnow()
            )
            
            # Save to database
            await self._save_document_info(document_info)
            
            # Store file in MongoDB GridFS
            await self._store_file_content(document_id, file_content, file.filename)
            
            # Queue for processing
            await self._queue_processing(document_id, temp_file_path)
            
            return document_info
            
        except Exception as e:
            logger.error(f"Error uploading document: {e}")
            raise HTTPException(status_code=500, detail="Failed to upload document")

    def _get_document_type(self, extension: str) -> DocumentType:
        """Determine document type from extension"""
        extension_map = {
            'pdf': DocumentType.PDF,
            'docx': DocumentType.DOCX,
            'doc': DocumentType.DOC,
            'txt': DocumentType.TXT,
            'html': DocumentType.HTML,
            'htm': DocumentType.HTML,
            'rtf': DocumentType.RTF,
            'jpg': DocumentType.IMAGE,
            'jpeg': DocumentType.IMAGE,
            'png': DocumentType.IMAGE,
            'tiff': DocumentType.IMAGE,
            'bmp': DocumentType.IMAGE
        }
        
        return extension_map.get(extension, DocumentType.TXT)

    async def _save_document_info(self, document_info: DocumentInfo):
        """Save document info to PostgreSQL"""
        query = """
            INSERT INTO documents (
                id, user_id, filename, original_filename, file_size,
                document_type, content_hash, processing_status,
                extracted_text, metadata
            ) VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, $10)
        """
        
        async with postgresql_manager.get_connection() as conn:
            await conn.execute(
                query, document_info.id, document_info.user_id,
                document_info.filename, document_info.original_filename,
                document_info.file_size, document_info.document_type,
                document_info.content_hash, document_info.processing_status,
                document_info.extracted_text, json.dumps(document_info.metadata)
            )

    async def _store_file_content(self, document_id: str, content: bytes, filename: str):
        """Store file content in MongoDB GridFS"""
        collection = mongodb_manager.get_collection("document_files")
        
        document = {
            "_id": document_id,
            "filename": filename,
            "content": content,
            "created_at": datetime.utcnow()
        }
        
        await collection.insert_one(document)

    async def _queue_processing(self, document_id: str, file_path: str):
        """Queue document for processing"""
        await hybrid_messaging_manager.publish_message(
            exchange="documents",
            routing_key="process.document",
            message={
                "document_id": document_id,
                "file_path": file_path,
                "timestamp": datetime.utcnow().isoformat()
            }
        )

    async def process_document(self, document_id: str, file_path: str) -> ProcessingResult:
        """Process document and extract content"""
        start_time = datetime.utcnow()
        
        try:
            # Update status to processing
            await self._update_processing_status(document_id, ProcessingStatus.PROCESSING)
            
            # Get document info
            document_info = await self.get_document(document_id)
            if not document_info:
                raise ValueError("Document not found")
            
            # Process based on document type
            processor = self.processors.get(document_info.document_type)
            if not processor:
                raise ValueError(f"No processor for document type: {document_info.document_type}")
            
            extracted_content = await processor.extract_content(file_path)
            
            # Update document with extracted content
            await self._update_extracted_content(document_id, extracted_content)
            
            # Update status to completed
            await self._update_processing_status(document_id, ProcessingStatus.COMPLETED)
            
            # Store in search index
            await self._index_document(document_id, extracted_content)
            
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            # Clean up temporary file
            try:
                os.remove(file_path)
            except Exception:
                pass
            
            return ProcessingResult(
                document_id=document_id,
                status=ProcessingStatus.COMPLETED,
                extracted_content=extracted_content.dict(),
                error_message=None,
                processing_time=processing_time
            )
            
        except Exception as e:
            logger.error(f"Error processing document {document_id}: {e}")
            
            # Update status to failed
            await self._update_processing_status(
                document_id, ProcessingStatus.FAILED, str(e)
            )
            
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            return ProcessingResult(
                document_id=document_id,
                status=ProcessingStatus.FAILED,
                extracted_content={},
                error_message=str(e),
                processing_time=processing_time
            )

    async def _update_processing_status(self, document_id: str, status: ProcessingStatus, error_message: str = None):
        """Update document processing status"""
        query = """
            UPDATE documents SET 
                processing_status = $1,
                error_message = $2,
                updated_at = NOW()
            WHERE id = $3
        """
        
        async with postgresql_manager.get_connection() as conn:
            await conn.execute(query, status, error_message, document_id)

    async def _update_extracted_content(self, document_id: str, content: ExtractedContent):
        """Update document with extracted content"""
        query = """
            UPDATE documents SET 
                extracted_text = $1,
                metadata = $2,
                updated_at = NOW()
            WHERE id = $3
        """
        
        async with postgresql_manager.get_connection() as conn:
            await conn.execute(
                query, content.text, json.dumps(content.dict()), document_id
            )

    async def _index_document(self, document_id: str, content: ExtractedContent):
        """Index document for search"""
        collection = mongodb_manager.get_collection("document_search")
        
        search_doc = {
            "_id": document_id,
            "text": content.text,
            "tables": content.tables,
            "images": content.images,
            "metadata": content.metadata,
            "indexed_at": datetime.utcnow()
        }
        
        await collection.replace_one(
            {"_id": document_id}, search_doc, upsert=True
        )

    async def get_document(self, document_id: str) -> Optional[DocumentInfo]:
        """Get document by ID"""
        query = "SELECT * FROM documents WHERE id = $1"
        
        async with postgresql_manager.get_connection() as conn:
            result = await conn.fetchrow(query, document_id)
            
            if result:
                data = dict(result)
                if data.get('metadata'):
                    data['metadata'] = json.loads(data['metadata'])
                return DocumentInfo(**data)
        
        return None

    async def get_user_documents(self, user_id: str) -> List[DocumentInfo]:
        """Get documents for user"""
        query = """
            SELECT * FROM documents 
            WHERE user_id = $1 
            ORDER BY created_at DESC
        """
        
        async with postgresql_manager.get_connection() as conn:
            results = await conn.fetch(query, user_id)
            
            documents = []
            for row in results:
                data = dict(row)
                if data.get('metadata'):
                    data['metadata'] = json.loads(data['metadata'])
                documents.append(DocumentInfo(**data))
            
            return documents

    async def search_documents(self, search_params: DocumentSearch) -> List[Dict[str, Any]]:
        """Search documents"""
        collection = mongodb_manager.get_collection("document_search")
        
        # Build search query
        query = {"$text": {"$search": search_params.query}}
        
        if search_params.user_id:
            # Need to join with PostgreSQL data - simplified for now
            pass
        
        cursor = collection.find(query).limit(search_params.limit)
        results = await cursor.to_list(length=search_params.limit)
        
        return results

    async def download_document(self, document_id: str) -> Optional[bytes]:
        """Download original document"""
        collection = mongodb_manager.get_collection("document_files")
        
        result = await collection.find_one({"_id": document_id})
        if result:
            return result["content"]
        
        return None

# Document processor consumer
class DocumentProcessor:
    def __init__(self, document_manager: DocumentManager):
        self.document_manager = document_manager
    
    async def start_consuming(self):
        """Start consuming processing messages"""
        await hybrid_messaging_manager.consume_messages(
            queue="document_processing_queue",
            callback=self.process_document_message
        )
    
    async def process_document_message(self, message: Dict[str, Any]):
        """Process document processing message"""
        try:
            document_id = message.get("document_id")
            file_path = message.get("file_path")
            
            if document_id and file_path:
                result = await self.document_manager.process_document(document_id, file_path)
                logger.info(f"Processed document {document_id}: {result.status}")
                
        except Exception as e:
            logger.error(f"Error processing document message: {e}")

# Initialize services
document_manager = DocumentManager()
document_processor = DocumentProcessor(document_manager)

# Lifespan context manager
@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    logger.info("Starting Document Service...")
    
    # Initialize database connections
    await postgresql_manager.initialize()
    await mongodb_manager.initialize()
    
    # Initialize messaging
    await hybrid_messaging_manager.initialize()
    
    # Create database tables
    await create_tables()
    
    # Start document processor
    asyncio.create_task(document_processor.start_consuming())
    
    logger.info("Document Service started successfully")
    
    yield
    
    # Shutdown
    logger.info("Shutting down Document Service...")
    await postgresql_manager.close()
    await mongodb_manager.close()
    await hybrid_messaging_manager.close()

# Create FastAPI app
app = FastAPI(
    title="Document Service",
    description="Document processing, content extraction, and search",
    version="1.0.0",
    lifespan=lifespan
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Database table creation
async def create_tables():
    """Create database tables"""
    queries = [
        """
        CREATE TABLE IF NOT EXISTS documents (
            id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            user_id UUID NOT NULL,
            filename VARCHAR(500) NOT NULL,
            original_filename VARCHAR(500) NOT NULL,
            file_size BIGINT NOT NULL,
            document_type VARCHAR(20) NOT NULL,
            content_hash VARCHAR(64) NOT NULL,
            processing_status VARCHAR(20) DEFAULT 'pending',
            extracted_text TEXT,
            metadata JSONB DEFAULT '{}',
            error_message TEXT,
            created_at TIMESTAMP DEFAULT NOW(),
            updated_at TIMESTAMP DEFAULT NOW()
        )
        """,
        "CREATE INDEX IF NOT EXISTS idx_documents_user_id ON documents(user_id);",
        "CREATE INDEX IF NOT EXISTS idx_documents_type ON documents(document_type);",
        "CREATE INDEX IF NOT EXISTS idx_documents_status ON documents(processing_status);",
        "CREATE INDEX IF NOT EXISTS idx_documents_hash ON documents(content_hash);"
    ]
    
    async with postgresql_manager.get_connection() as conn:
        for query in queries:
            await conn.execute(query)

# API Routes
@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "service": "document-service"}

@app.post("/upload", response_model=DocumentInfo)
async def upload_document(
    user_id: str,
    file: UploadFile = File(...),
    background_tasks: BackgroundTasks = None
):
    """Upload document"""
    if file.size > 50 * 1024 * 1024:  # 50MB limit
        raise HTTPException(status_code=413, detail="File too large")
    
    return await document_manager.upload_document(user_id, file)

@app.get("/documents/{document_id}", response_model=DocumentInfo)
async def get_document(document_id: str):
    """Get document by ID"""
    document = await document_manager.get_document(document_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    return document

@app.get("/user/{user_id}/documents", response_model=List[DocumentInfo])
async def get_user_documents(user_id: str):
    """Get user documents"""
    return await document_manager.get_user_documents(user_id)

@app.post("/search")
async def search_documents(search_params: DocumentSearch):
    """Search documents"""
    return await document_manager.search_documents(search_params)

@app.get("/download/{document_id}")
async def download_document(document_id: str):
    """Download document"""
    content = await document_manager.download_document(document_id)
    if not content:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Get document info for filename
    doc_info = await document_manager.get_document(document_id)
    filename = doc_info.original_filename if doc_info else "document"
    
    return StreamingResponse(
        io.BytesIO(content),
        media_type='application/octet-stream',
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8006) 